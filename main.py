import os
import re
import random
import html  # بالای فایل
from io import BytesIO
from telegram.helpers import escape
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.constants import ParseMode
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters
)
from supabase import create_client, Client
from docx import Document
from dotenv import load_dotenv
       



load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
OWNER_ID = int(os.getenv("OWNER_ID"))



supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

app = ApplicationBuilder().token(BOT_TOKEN).build()

user_states = {}  # اضافه شد
quiz_sessions = {}

CATEGORIES = ["Nomen", "Verb", "Adjektiv", "Adverb"]

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return
    await update.message.reply_text("📄 لطفاً کلمه خود را وارد کنید")
    user_states[update.effective_user.id] = {"step": "word"}


async def message_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    user_id = update.effective_user.id

    # 👇👇 بررسی اینکه کاربر در حالت افزودن جمله هست
    if "add_example_word" in context.user_data:
        word_data = context.user_data.pop("add_example_word")
        word_id = word_data["index"]
        user_id_str = str(user_id)
        new_example = update.message.text.strip()

        result = supabase.from_("words").select("*") \
            .eq("user_id", user_id_str).eq("index", word_id).execute().data
        if result:
            current = result[0]
            examples = current.get("examples") or []
            examples.append(new_example)

            supabase.from_("words").update({"examples": examples}) \
                .eq("user_id", user_id_str).eq("index", word_id).execute()

            await update.message.reply_text("✅ جمله با موفقیت ذخیره شد.")
        return  # مهم! چون نمی‌خوای بقیه کد اجرا شه

    # ✅ ادامه کد مربوط به /start:
    user_id = update.effective_user.id
    state = user_states.get(user_id)
    if not state:
        return

    text = update.message.text.strip()
    if state["step"] == "word":
        state["word"] = text
        state["step"] = "meaning"
        await update.message.reply_text("🧠 حالا معنی کلمه را وارد کنید")
    elif state["step"] == "meaning":
        state["meaning"] = text
        state["step"] = "category"
        buttons = [[InlineKeyboardButton(cat, callback_data=f"category:{cat}")] for cat in CATEGORIES]
        await update.message.reply_text("📂 لطفاً دسته‌بندی را انتخاب کن:", reply_markup=InlineKeyboardMarkup(buttons))



def save_word(user_id, word, meaning, category):
    result = supabase.from_("words") \
        .select("index") \
        .eq("user_id", str(user_id)) \
        .order("index", desc=True) \
        .limit(1) \
        .execute()
    last_index = result.data[0]["index"] if result.data else 0
    new_index = last_index + 1

    supabase.from_("words").insert({
        "word": word,
        "meaning": meaning,
        "category": category,
        "user_id": str(user_id),
        "index": new_index
    }).execute()

import html  # مطمئن شو این بالای فایل هست

async def list_words(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    try:
        user_id_str = str(update.effective_user.id)
        args = context.args

        if args:
            selected_category = args[0].capitalize()
            if selected_category not in CATEGORIES:
                await update.message.reply_text("❗ دسته‌بندی معتبر نیست. دسته‌های مجاز: Nomen, Verb, Adjektiv, Adverb")
                return
            words = supabase.from_("words").select("*") \
                .eq("user_id", user_id_str) \
                .eq("category", selected_category) \
                .order("index").execute().data
        else:
            words = supabase.from_("words").select("*") \
                .eq("user_id", user_id_str) \
                .order("index").execute().data

        if not words:
            await update.message.reply_text("⚠️ هیچ کلمه‌ای ذخیره نشده.")
            return

        text = "📚 <b>کلمه‌های ذخیره‌شده:</b>\n\n"
        messages = []
        for w in words:
            index = w.get("index", "-")
            word = html.escape(w.get("word", "❓"))
            meaning = html.escape(w.get("meaning", "❓"))
            category = html.escape(w.get("category", "❓بدون دسته‌بندی"))
            examples = w.get("examples") or []

            block = f"{index}. <b>{word}</b> ➜ {meaning} ({category})\n"
            for ex in examples:
                block += f"📝 {html.escape(ex)}\n"
            block += "\n"

            if len(text + block) > 4000:
                messages.append(text)
                text = ""

            text += block

        if text:
            messages.append(text)

        for msg in messages:
            await update.message.reply_text(msg.strip(), parse_mode=ParseMode.HTML)

    except Exception as e:
        print(f"❌ خطای غیرمنتظره در list_words: {e}")
        await update.message.reply_text("🚫 خطایی در اجرای دستور پیش آمد.")


    

async def export_all(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    words = supabase.from_("words").select("*").eq("user_id", str(update.effective_user.id)).order("index").execute().data
    if not words:
        await update.message.reply_text("⚠️ هیچ کلمه‌ای برای خروجی وجود ندارد.")
        return

    doc = Document()
    doc.add_heading("تمام کلمات ذخیره‌شده", 0)
    for item in words:
        doc.add_heading(f"{item['index']}. {item['word']}", level=1)
        doc.add_paragraph(f"🔹 معنی: {item['meaning']}")
        doc.add_paragraph(f"🏷 دسته‌بندی: {item.get('category', 'بدون دسته‌بندی')}")
        examples = item.get("examples") or []
        if examples:
            doc.add_paragraph("📝 مثال‌ها:")
            for ex in examples:
                doc.add_paragraph(f"• {ex}", style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    await update.message.reply_document(document=buffer, filename="alle_woerter.docx")

async def button_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    state = user_states.get(user_id)
    if not state or not query.data.startswith("category:"):
        return

    category = query.data.split(":")[1]
    word = state["word"]
    meaning = state["meaning"]

    try:
        save_word(user_id, word, meaning, category)
        await query.edit_message_text(f"✅ کلمه '{word}' با موفقیت در دسته‌بندی '{category}' ذخیره شد.")
    except Exception as e:
        await query.edit_message_text(f"❌ خطا در ذخیره‌سازی:{e}")

    user_states.pop(user_id)


async def add_example_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    if not context.args:
        await update.message.reply_text("❗ لطفاً کلمه یا شماره را وارد کن: /addexample Haus یا /addexample 3")
        return

    keyword = " ".join(context.args).strip()
    words = supabase.from_("words").select("*").eq("user_id", str(update.effective_user.id)).execute().data
    selected = None
    if keyword.isdigit():
        selected = next((w for w in words if w.get("index") == int(keyword)), None)
    else:
        selected = next((w for w in words if w.get("word") == keyword), None)

    if not selected:
        await update.message.reply_text("❌ کلمه‌ای پیدا نشد.")
        return

    context.user_data["add_example_word"] = selected
    await update.message.reply_text(f'✍ لطفاً جمله‌ای برای "{selected["word"]}" ارسال کنید:')



async def export_words(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    if not context.args:
        await update.message.reply_text("❗ لطفاً شماره‌ها را بعد از /export بنویس، مثل: /export 1 3 5")
        return

    try:
        indexes = [int(x) for x in context.args if x.isdigit()]
    except ValueError:
        await update.message.reply_text("❌ لطفاً فقط شماره‌های معتبر وارد کن.")
        return

    data = supabase.from_("words").select("*").eq("user_id", str(update.effective_user.id)).execute().data
    filtered = [w for w in data if w.get("index") in indexes]

    if not filtered:
        await update.message.reply_text("❌ کلمه‌ای با این شماره‌ها پیدا نشد.")
        return

    text = "📋 <b>کلمه‌های انتخاب‌شده:</b>\n\n"
    for w in filtered:
        text += f"{w['index']}. <b>{w['word']}</b> ➜ {w['meaning']}\n"
        examples = w.get("examples") or []
        if examples:
            for ex in examples:
                text += f"📝 {ex}\n"
        text += "\n"

    await update.message.reply_text(text.strip(), parse_mode=ParseMode.HTML)

    # ساخت فایل Word
    doc = Document()
    doc.add_heading("Exportierte Wörter", 0)
    for item in filtered:
        doc.add_heading(f"{item['index']}. {item['word']}", level=1)
        doc.add_paragraph(f"🔹 معنی: {item['meaning']}")
        examples = item.get("examples") or []
        if examples:
            doc.add_paragraph("📝 مثال‌ها:")
            for ex in examples:
                doc.add_paragraph(f"• {ex}", style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    await update.message.reply_document(document=buffer, filename="woerter_export.docx")


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return
    text = (
        "📌 <b>دستورات موجود:</b>\n\n"
        "/start – افزودن کلمه جدید\n"
        "/list – نمایش همه کلمات\n"
        "/quiz – آزمون چهارگزینه‌ای\n"
        "/addexample – افزودن جمله\n"
        "/export - خروجی گرفتن\n"
        "/showall -نمایش همه کلمات \n"
        "/help – نمایش همین راهنما"
    )
    await update.message.reply_text(text, parse_mode=ParseMode.HTML)

async def quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return
    data = supabase.from_("words").select("*").eq("user_id", str(update.effective_user.id)).execute().data
    if len(data) < 4:
        await update.message.reply_text("⚠️ حداقل ۴ کلمه لازم است.")
        return

    session = {
        "items": random.sample(data, min(10, len(data))),
        "current": 0,
        "score": 0
    }
    context.user_data["quiz"] = session
    await ask_question(update, context)

async def ask_question(update: Update, context: ContextTypes.DEFAULT_TYPE):
    session = context.user_data["quiz"]
    item = session["items"][session["current"]]
    options = random.sample(session["items"], 3)
    options.append(item)
    random.shuffle(options)

    buttons = [[InlineKeyboardButton(o["meaning"], callback_data=o["word"])] for o in options]
    await update.message.reply_text(
        f"❓ سوال {session['current'] + 1} از {len(session['items'])}: <b>{item['word']}</b> یعنی چی؟",
        parse_mode=ParseMode.HTML,
        reply_markup=InlineKeyboardMarkup(buttons)
    )

async def answer_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    session = context.user_data.get("quiz")
    if not session:
        return
    current_item = session["items"][session["current"]]
    if query.data == current_item["word"]:
        session["score"] += 1
        await query.edit_message_text("✅ آفرین درست گفتی")
    else:
        await query.edit_message_text(f"❌ جواب اشتباه بود. معنی درست: {current_item['meaning']}")

    session["current"] += 1
    if session["current"] < len(session["items"]):
        await ask_question(query, context)
    else:
        await context.bot.send_message(
            chat_id=query.from_user.id,
            text=f"🏁 آزمون تمام شد. امتیاز: {session['score']} از {len(session['items'])}"
        )
async def show_all_words(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_user.id != OWNER_ID:
        return

    words = supabase.from_("words").select("*") \
        .eq("user_id", str(update.effective_user.id)) \
        .order("index").execute().data

    if not words:
        await update.message.reply_text("⚠️ هیچ کلمه‌ای ذخیره نشده.")
        return

    MAX_MESSAGE_LENGTH = 4000  # محدودیت تلگرام
    message_parts = []
    current_text = ""

    for w in words:
        index = w.get("index", "❓")
        word = html.escape(w.get("word", "❓"))
        meaning = html.escape(w.get("meaning", "❓"))
        examples = w.get("examples") or []

        block = f"{index}. <b>{word}</b> ➜ {meaning}\n"
        for ex in examples:
            block += f"📝 {html.escape(ex)}\n"
        block += "\n"

        if len(current_text) + len(block) > MAX_MESSAGE_LENGTH:
            message_parts.append(current_text)
            current_text = block
        else:
            current_text += block

    if current_text:
        message_parts.append(current_text)

    for part in message_parts:
        await update.message.reply_text(part.strip(), parse_mode=ParseMode.HTML)


app.add_handler(CommandHandler("start", start))
app.add_handler(CommandHandler("list", list_words))
app.add_handler(CommandHandler("quiz", quiz))
app.add_handler(CommandHandler("export", export_words))
app.add_handler(CommandHandler("exportall", export_all))
app.add_handler(CommandHandler("addexample", add_example_command))
app.add_handler(CommandHandler("help", help_command))
app.add_handler(CommandHandler("showall", show_all_words))

app.add_handler(CallbackQueryHandler(button_handler, pattern="^category:.*$"))
app.add_handler(CallbackQueryHandler(answer_callback))

app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, message_handler))


import asyncio

PORT = int(os.environ.get("PORT", "8080"))
RENDER_HOST = os.environ.get("RENDER_EXTERNAL_HOSTNAME")
WEBHOOK_PATH = f"/{BOT_TOKEN}"
WEBHOOK_URL = f"https://{RENDER_HOST}/{BOT_TOKEN}"

if __name__ == "__main__":
    asyncio.run(app.run_webhook(
        listen="0.0.0.0",
        port=PORT,
        url_path=WEBHOOK_PATH,
        webhook_url=WEBHOOK_URL
    ))

